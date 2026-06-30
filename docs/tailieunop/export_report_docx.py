"""Convert TECHNICAL_REPORT_EN.md to a formatted .docx file."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
MD_FILE = ROOT / "TECHNICAL_REPORT_EN.md"
DOCX_FILE = ROOT / "TECHNICAL_REPORT.docx"


def set_document_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(6)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(col_count):
            cell = table.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[-:\s|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    set_document_styles(doc)

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
            i += 1
            continue
        if stripped.startswith("##### "):
            doc.add_heading(stripped[6:], level=4)
            i += 1
            continue

        if stripped.startswith("- [ ] "):
            doc.add_paragraph(stripped[6:], style="List Bullet")
            i += 1
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            i += 1
            continue

        # Bold label lines like **Course:**
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    if not MD_FILE.exists():
        raise SystemExit(f"Missing {MD_FILE}")
    md_to_docx(MD_FILE, DOCX_FILE)
